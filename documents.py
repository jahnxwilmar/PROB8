import os
import sys
import comtypes.client
import calendar
from datetime import datetime
from docx import Document
from models import get_barangay_captain, get_barangay_kagawad, get_barangay_treasurer, get_barangay_secretary, get_barangay_pangkat_member, get_barangay_sk
import webbrowser
from pathlib import Path
from tkinter import messagebox
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, CT_Tbl
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.table import Table
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document as f_Document
from docx.shared import Inches


# from docx2pdf import convert
# import sys


APP_NAME = "PROB8"


str_dayToday = int(datetime.now().day)
str_yearToday = str(datetime.now().year)
str_yearFromToday = str(datetime.now().year + 1)
month_number = datetime.today().month
month_name = calendar.month_name[month_number]
pre_date_today = datetime.today()
post_date_today = pre_date_today.strftime("%m-%d-%Y")


def get_unique_filename(filepath):
    counter = 1
    filename, extension = os.path.splitext(filepath)
    
    while os.path.exists(filepath):
        filepath = f"{filename} ({counter}){extension}"
        counter += 1
        
    return filepath


def replace_with_table(doc: f_Document, elm, parent, rows: int, cols: int) -> Table:
    table_elm = CT_Tbl.new_tbl(rows, cols, doc._block_width)
    elm.getparent().replace(elm, table_elm)
    table = Table(table_elm, parent)

    # Center the table in the document
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblPr.append(OxmlElement('w:jc'))
    tblPr.xpath('w:jc')[0].set(qn('w:val'), 'center')

    # Add borders to the table
    for cell in table._cells:
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border in ['top', 'left', 'bottom', 'right']:
            border_elm = OxmlElement(f'w:{border}')
            border_elm.set(qn('w:val'), 'single')
            border_elm.set(qn('w:sz'), '4')  # Border size
            border_elm.set(qn('w:space'), '0')
            border_elm.set(qn('w:color'), '000000')  # Black color
            tcBorders.append(border_elm)
        tcPr.append(tcBorders)

    return table


def set_cell_border(cell, **kwargs):
    """
    Set borders for a cell.
    Usage:
    set_cell_border(cell, bottom={"sz": 12, "val": "single", "color": "000000"})
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    for border_name in ['top', 'start', 'end', 'bottom', 'insideH', 'insideV']:
        if border_name in kwargs:
            border = OxmlElement(f'w:{border_name}')
            for key in kwargs[border_name]:
                border.set(qn(f'w:{key}'), kwargs[border_name][key])
            tcPr.append(border)

def center_text_in_cell(cell):
    """
    Center align the text in the cell.
    """
    # Clear any existing alignment
    p = cell.paragraphs[0]
    p.alignment = 1  # 1 corresponds to CENTER alignment in python-docx

# Load the document

def test_center_text_in_cell(cell, font_size=6):
    """
    Center align the text in the cell and set font size.
    """
    # Clear any existing alignment and formatting
    p = cell.paragraphs[0]
    p.alignment = 1  # 1 corresponds to CENTER alignment in python-docx
    
    # Apply font size
    run = p.runs[0] if p.runs else p.add_run()
    run.font.size = Pt(font_size)

# Load the document

def set_tables(doc, respondents, complainants, tab):
    # Locate the table (assuming it's the first table)
    table = doc.tables[-tab]  # Adjust the index to select the correct table

    # Remove all rows from the table
    for row in table.rows:
        tbl = row._element
        tbl.getparent().remove(tbl)

    data = []

    list_range = 0
    if len(respondents) > len(complainants):
        list_range = len(respondents)
    else:
        list_range = len(complainants)

    for i in range(list_range):
        first_row = ""
        second_row = ""
        third_row = ""

        try:
            value = respondents[i]
        except IndexError:
            first_row = ""
        else:
            first_row = value

        try:
            value = complainants[i]
        except IndexError:
            third_row = ""
        else:
            third_row = value


        data.append([first_row, second_row, third_row])

    # Add data to the table and set bottom borders
    for row_data in data:
        row = table.add_row()  # Add a new row
        counter = 1
        for idx, cell_data in enumerate(row_data):
            if counter != 2:
                cell = row.cells[idx]
                cell.text = cell_data
                set_cell_border(cell, bottom={"sz": "4", "val": "single", "color": "000000"})
                center_text_in_cell(cell)
            counter += 1
        counter = 1


def set_report_tables(doc, values, tab):
    # Locate the table (assuming it's the first table)
    table = doc.tables[-tab]  # Adjust the index to select the correct table

    # Remove all rows from the table
    for row in table.rows:
        tbl = row._element
        tbl.getparent().remove(tbl)

    data = []

    for i in range(len(values)):
        first_row = ""
        second_row = ""
        third_row = ""
        fourth_row = ""
        fifth_row = ""
        sixth_row = ""

        first_row = str(i + 1)
        second_row = values[i][0]
        third_row = values[i][1]
        fourth_row = values[i][2]
        fifth_row = values[i][3]
        sixth_row = values[i][4]

        data.append((first_row, second_row, third_row, fourth_row, fifth_row, sixth_row))

    for row_data in data:
        row = table.add_row()  # Add a new row
        set_row_height(row, 500)  # Set row height (500 twips = 25 points)
        
        counter = 1
        for idx, cell_data in enumerate(row_data):
            if counter != 2:
                cell = row.cells[idx]
                cell.text = cell_data
                center_text_in_cell(cell)
            else:
                cell = row.cells[idx]
                cell.text = f"     " + cell_data
            counter += 1
        counter = 1

    # Add data to the table and set bottom borders
    # for row_data in data:
    #     row = table.add_row()  # Add a new row
    #     counter = 1
    #     for idx, cell_data in enumerate(row_data):
    #         if counter != 10:
    #             cell = row.cells[idx]
    #             cell.text = cell_data
    #             center_text_in_cell(cell)
    #         counter += 1
    #     counter = 1


def set_row_height(row, height):
    """
    Set the height of the row in twips (twentieth of a point).
    """
    tr = row._tr  # Get the row's XML element
    trPr = tr.get_or_add_trPr()  # Get or create trPr element
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))  # Set the height
    trHeight.set(qn('w:hRule'), 'exact')  # Use 'exact' to ensure the height is applied exactly
    trPr.append(trHeight)


def center_text_in_cell(cell):
    """Center text in the cell"""
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def resources_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)


def get_resident_name(firstname, middlename, lastname, suffix):
    fullname = f"{firstname} {middlename[0]}. {lastname}"
    if not suffix == "":
        fullname += f" {suffix}"
    return fullname


def add_extended_paragraphs(doc, old_text, respondent_name):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    parts = inline[i].text.split(old_text)
                    
                    inline[i].text = parts[0]

                    day_suffix = get_ordinal_suffix(str_dayToday)

                    for i in range(2, len(respondent_name)):
                        new_run_1 = p.add_run(" and upon respondent ")
                        new_run_2 = p.add_run(respondent_name[i].upper())
                        new_run_2.font.underline = True
                        new_run_3 = p.add_run(" on the ")
                        new_run_4 = p.add_run(str(str_dayToday))
                        new_run_5 = p.add_run(day_suffix)
                        new_run_5.font.superscript = True
                        new_run_4.font.underline = True
                        new_run_5.font.underline = True
                        new_run_6 = p.add_run(" day of ")
                        new_run_7 = p.add_run(month_name)
                        new_run_7.font.underline = True
                        new_run_8 = p.add_run(f", {str(str_yearToday)}")

                    final_new_run = p.add_run(" by:")

                    if len(parts) > 1:
                        p.add_run(parts[1])


def add_counterclaim_extended_paragraphs(doc, old_text, respondent_name):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    parts = inline[i].text.split(old_text)
                    
                    inline[i].text = parts[0]

                    day_suffix = get_ordinal_suffix(str_dayToday)

                    counter = 1
                    for i in range(2, len(respondent_name)):
                        if counter == 1:
                            new_run_1 = p.add_run("and ")
                        else:
                            new_run_1 = p.add_run(" and ")
                        new_run_2 = p.add_run(respondent_name[i].upper())
                        new_run_2.font.underline = True
                        new_run_3 = p.add_run("(name)")

                        counter += 1

                    final_new_run = p.add_run(" have been found to have willfully failed or refused to appear without justifiable reason before the Punong Barangay/Pangkat ng Tagapagkasundo and therefore respondent/s is are barred from filing his/their counterclaim (if any) arising from the complaint in court/government office.")

                    if len(parts) > 1:
                        p.add_run(parts[1])


def add_action_extended_paragraphs(doc, old_text, respondent_name, date):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    parts = inline[i].text.split(old_text)
                    
                    inline[i].text = parts[0]

                    day_suffix = get_ordinal_suffix(str_dayToday)

                    pre_run_1 = p.add_run("This is to certify that the above-captioned case was dismissed pursuant to the Order dated ")

                    pre_run_2 = p.add_run(date)
                    pre_run_2.font.underline = True

                    pre_run_3 = p.add_run(", for complainant/s ")

                    pre_run_4 = p.add_run(respondent_name[0].upper())
                    pre_run_4.font.underline = True

                    pre_run_5 = p.add_run("(name) and ")

                    if len(respondent_name) > 1:
                        pre_run_6 = p.add_run(respondent_name[1].upper())
                        pre_run_6.font.underline = True
                    else:
                        pre_run_6 = p.add_run("________________________________")

                    pre_run_7 = p.add_run("(name) ")

                    if len(respondent_name) > 2:
                        counter = 1
                        for i in range(2, len(respondent_name)):
                            if counter == 1:
                                new_run_1 = p.add_run("and ")
                            else:
                                new_run_1 = p.add_run(" and ")
                            new_run_2 = p.add_run(respondent_name[i].upper())
                            new_run_2.font.underline = True
                            new_run_3 = p.add_run("(name)")

                            counter += 1

                    final_new_run = p.add_run("willful failure or refusal to appear for hearing before the Punong Barangay/Pangkat ng Tagapagkasundo and therefore complainant/s is/are barred from filing an action in court/government office.")

                    if len(parts) > 1:
                        p.add_run(parts[1])


def set_report_filters(doc, filter_values):
    # Locate the table (assuming it's the first table)
    table = doc.tables[0]  # Adjust the index to select the correct table

    # Remove all rows from the table
    for row in table.rows:
        tbl = row._element
        tbl.getparent().remove(tbl)

    data = []

    temp_residential_status = filter_values['resident_status'].split(', ')

    if len(filter_values) == 0:
        data.append((f"- All the residents of BARANGAY POBLACION 8"))
    else:
        if len(filter_values) == 1 and len(temp_residential_status) == 1 and temp_residential_status[0] == 'Active':
            data.append((f"- All the residents of BARANGAY POBLACION 8"))
        elif len(filter_values) == 1 and len(temp_residential_status) == 1 and temp_residential_status[0] == 'Transferred':
            data.append((f"- All the transferred residents of BARANGAY POBLACION 8"))
        elif len(filter_values) == 1 and len(temp_residential_status) == 1 and temp_residential_status[0] == 'Deceased':
            data.append((f"- All the deceased residents of BARANGAY POBLACION 8"))
        elif len(filter_values) == 1 and len(temp_residential_status) == 1 and temp_residential_status[0] == 'Transient':
            data.append((f"- All the transient residents of BARANGAY POBLACION 8"))
        elif len(temp_residential_status) > 1:
            data.append(f"- RESIDENTIAL STATUS: {filter_values['resident_status']}")

        if 'purok' in filter_values:
            data.append((f"- PUROK: {filter_values['purok']}"))
        if 'age' in filter_values:
            age_value = filter_values['age'].split(', ')
            if age_value[0] == age_value[1]:
                data.append(f"- AGE: {age_value[0]}")
            else:
                data.append(f"- AGE: {age_value[0]} to {age_value[1]}")
        if 'sex' in filter_values:
            data.append(f"- SEX: {filter_values['sex']}")
        if 'comelec' in filter_values:
            data.append(f"- COMELEC: {filter_values['comelec']}")
        if 'philsys' in filter_values:
            data.append(f"- PHILSYS: {filter_values['philsys']}")
        if 'civil_status' in filter_values:
            data.append((f"- CIVIL STATUS: {filter_values['civil_status']}"))
        if 'blood_type' in filter_values:
            data.append((f"- BLOOD TYPE: {filter_values['blood_type']}"))
        if 'educational_status' in filter_values:
            data.append((f"- EDUCATIONAL STATUS: {filter_values['educational_status']}"))
        if 'memberships_orgs' in filter_values:
            data.append((f"- MEMBERSHIP/ORGANIZATIONS: {filter_values['memberships_orgs']}"))

    for row_data in data:
        row = table.add_row()  # Add a new row

        cell = row.cells[0]
        cell.text = row_data


def add_reports_paragraphs(doc, old_text, filter_values):
    if len(filter_values) == 0:
        for p in doc.paragraphs:
            if old_text in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        parts = inline[i].text.split(old_text)
                        
                        inline[i].text = parts[0]

                        new_run = p.add_run("        - All the residents of Poblacion 8.")

                        if len(parts) > 1:
                            p.add_run(parts[1])
    else:
        for p in doc.paragraphs:
            if old_text in p.text:
                # Split the current paragraph where old_text is found
                inline = p.runs
                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        parts = inline[i].text.split(old_text)
                        
                        # Replace the text before the placeholder
                        inline[i].text = parts[0]

                        if 'purok' in filter_values:
                            new_para = p.insert_paragraph_before()
                            new_run = new_para.add_run("        - PUROK: ")
                            new_run.font.bold = True
                            new_run = new_para.add_run(str(filter_values['purok']))

                        if 'age' in filter_values:
                            age_value = filter_values['age'].split(',')

                            new_para = p.insert_paragraph_before()
                            new_run = new_para.add_run("        - AGE: ")
                            new_run.font.bold = True
                            new_run = new_para.add_run(f"{age_value[0]} to {age_value[1]}")

                        # Add the rest of the original text (after the placeholder)
                        if len(parts) > 1:
                            p.add_run(parts[1])

                        return


def replace_text_in_paragraphs_underline(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    # Split the run's text at the old text
                    parts = inline[i].text.split(old_text)
                    
                    # Rebuild the run with the text before the old text, the new underlined text, and the text after
                    inline[i].text = parts[0]
                    
                    # Add a new run for the new underlined text
                    new_run = p.add_run(new_text)
                    new_run.font.underline = True
                    
                    # Add the text after the old text as a new run
                    if len(parts) > 1:
                        p.add_run(parts[1])

    # for paragraph in doc.paragraphs:
    #     if placeholder in paragraph.text:
    #         # Iterate over the runs within the paragraph
    #         for run in paragraph.runs:
    #             if placeholder in run.text:
    #                 # Find the index of the word within the run's text
    #                 start_index = run.text.find(placeholder)
    #                 end_index = start_index + len(placeholder)

    #                 # Split the text into before, target word, and after
    #                 before = run.text[:start_index]
    #                 target_word = run.text[start_index:end_index]
    #                 after = run.text[end_index:]

    #                 # Update the run with the text before the word
    #                 run.text = before

    #                 # Create a new run for the underlined word and add it after the original run
    #                 underlined_run = paragraph.add_run(target_word)
    #                 underlined_run.font.underline = True

    #                 # Add a new run for the text after the word
    #                 paragraph.add_run(after)


def replace_text_in_paragraphs(paragraphs, old_text, new_text):
    for p in paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text


def docx_find_replace_text(doc, search_text, replacing_text):
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        if search_text in p.text:
            inline = p.runs
            # Replace strings and retain the same style.
            # The text to be replaced can be split over several runs so
            # search through, identify which runs need to have text replaced
            # then replace the text in those identified
            started = False
            search_index = 0
            # found_runs is a list of (inline index, index of match, length of match)
            found_runs = list()
            found_all = False
            replace_done = False
            for i in range(len(inline)):

                # case 1: found in single run so short circuit the replacement
                if search_text in inline[i].text and not started:
                    found_runs.append((i, inline[i].text.find(search_text), len(search_text)))
                    text = inline[i].text.replace(search_text, str(replacing_text))
                    inline[i].text = text
                    replace_done = True
                    found_all = True
                    break

                if search_text[search_index] not in inline[i].text and not started:
                    # keep looking ...
                    continue

                # case 2: search for partial text, find first run
                if search_text[search_index] in inline[i].text and inline[i].text[-1] in search_text and not started:
                    # check sequence
                    start_index = inline[i].text.find(search_text[search_index])
                    check_length = len(inline[i].text)
                    for text_index in range(start_index, check_length):
                        if inline[i].text[text_index] != search_text[search_index]:
                            # no match so must be false positive
                            break
                    if search_index == 0:
                        started = True
                    chars_found = check_length - start_index
                    search_index += chars_found
                    found_runs.append((i, start_index, chars_found))
                    if search_index != len(search_text):
                        continue
                    else:
                        # found all chars in search_text
                        found_all = True
                        break

                # case 2: search for partial text, find subsequent run
                if search_text[search_index] in inline[i].text and started and not found_all:
                    # check sequence
                    chars_found = 0
                    check_length = len(inline[i].text)
                    for text_index in range(0, check_length):
                        if inline[i].text[text_index] == search_text[search_index]:
                            search_index += 1
                            chars_found += 1
                        else:
                            break
                    # no match so must be ended
                    found_runs.append((i, 0, chars_found))
                    if search_index == len(search_text):
                        found_all = True
                        break

            if found_all and not replace_done:
                for i, item in enumerate(found_runs):
                    index, start, length = [t for t in item]
                    if i == 0:
                        text = inline[index].text.replace(inline[index].text[start:start + length], str(replacing_text))
                        inline[index].text = text
                    else:
                        text = inline[index].text.replace(inline[index].text[start:start + length], '')
                        inline[index].text = text


# def convert_docx_to_pdf(docx_path, pdf_path):
#     try:
#         sys.stderr = open("logs/console-output.log", "w")
#         convert(docx_path, pdf_path)
#     except Exception as e:
#         print(f"An error occurred: {e}")
#     finally:
#         sys.stderr.close()
#         sys.stderr = sys.__stderr__  # Reset stderr to default
#
#         # Open the generated PDF in the web browser
#     webbrowser.open(pdf_path)


def convert_pdf(input_file_path, output_file_path):

    wdFormatPDF = 17

    in_file = input_file_path
    out_file = output_file_path
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False  # Ensure Word remains hidden
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

    webbrowser.open(output_file_path)


def open_word_file(file_path):
    try:
        os.startfile(file_path)
    except Exception as e:
        print(f"An error occurred while opening the file: {e}")


def replace_text(doc, old_text, new_text):
    replace_text_in_paragraphs(doc.paragraphs, old_text, new_text)
    docx_find_replace_text(doc, old_text, new_text)


def replace_text_version_1(doc, old_text, new_text):
    docx_find_replace_text(doc, old_text, new_text)


def replace_text_version_2(doc, old_text, new_text):
    replace_text_in_paragraphs(doc.paragraphs, old_text, new_text)


def get_ordinal_suffix(number):
    if 10 <= number % 100 <= 20:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(number % 10, 'th')
    return f"{suffix}"


def do_certificateLowIncome(fullName, purok, gender, reason, captain):
    doc_path = resources_path('assets\\templates\\Certificate_LowIncome.docx')  # Adjusted path
    doc = Document(doc_path)

    personal_pronoun = ""
    pronoun = ""

    if gender.lower() == 'male':
        personal_pronoun = "him"
        pronoun = "he"
    elif gender.lower() == 'female':
        personal_pronoun = "her"
        pronoun = "she"

    day_suffix = get_ordinal_suffix(str_dayToday)

    replace_text_version_2(doc, 'RESIDENTNAME', fullName.upper())
    replace_text_version_2(doc, 'RESIDENTPUROK', purok.capitalize())
    replace_text_version_2(doc, 'PERSONALPRONOUN', personal_pronoun)
    replace_text_version_1(doc, 'RESIDENTREASON', reason)
    replace_text_version_1(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_1(doc, 'DAYSUFFIX', day_suffix)
    replace_text_version_1(doc, 'MONTHTODAY', month_name)
    replace_text_version_1(doc, 'YEARTODAY', str(datetime.now().year))
    replace_text_version_1(doc, 'KAPITANNAME', captain.upper())

    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Certificate of Low Income")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Certificate of Low Income")
    document_name = f"{fullName} - Certificate of Low Income.docx"
    pdf_name = f"{fullName} - Certificate of Low Income.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    # Save the document
    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_certificateLowIncome_on_duty(fullName, purok, gender, reason, captain, officer):
    doc_path = resources_path('assets\\templates\\Certificate_LowIncome_on_duty.docx')  # Adjusted path
    doc = Document(doc_path)

    personal_pronoun = ""
    pronoun = ""

    if gender.lower() == 'male':
        personal_pronoun = "him"
        pronoun = "he"
    elif gender.lower() == 'female':
        personal_pronoun = "her"
        pronoun = "she"

    day_suffix = get_ordinal_suffix(str_dayToday)

    replace_text_version_2(doc, 'RESIDENTNAME', fullName.upper())
    replace_text_version_2(doc, 'RESIDENTPUROK', purok.capitalize())
    replace_text_version_2(doc, 'PERSONALPRONOUN', personal_pronoun)
    replace_text_version_1(doc, 'RESIDENTREASON', reason)
    replace_text_version_1(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_1(doc, 'DAYSUFFIX', day_suffix)
    replace_text_version_1(doc, 'MONTHTODAY', month_name)
    replace_text_version_1(doc, 'YEARTODAY', str(datetime.now().year))
    replace_text_version_1(doc, 'KAPITANNAME', captain.upper())
    replace_text_version_1(doc, 'KGWDNAME', officer.upper())

    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Certificate of Low Income")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Certificate of Low Income")
    document_name = f"{fullName} - Certificate of Low Income.docx"
    pdf_name = f"{fullName} - Certificate of Low Income.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    # Save the document
    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_certificateOfGoodMoral(fullName, purok, gender, civil, reason, ctc, ctc_date, captain):
    doc_path = resources_path('assets\\templates\\Certificate_GoodMoral.docx')  # Adjusted path
    doc = Document(doc_path)

    personal_pronoun = ""
    pronoun = ""

    if gender.lower() == 'male':
        personal_pronoun = "him"
        pronoun = "he"
    elif gender.lower() == 'female':
        personal_pronoun = "her"
        pronoun = "she"

    day_suffix = get_ordinal_suffix(str_dayToday)

    replace_text_version_2(doc, 'RESIDENTNAME', fullName.upper())
    replace_text_version_2(doc, 'RESIDENTPUROK', purok.capitalize())
    replace_text_version_2(doc, 'PSRONOUN', pronoun)
    replace_text_version_2(doc, 'PERSONALPRONOUN', personal_pronoun)
    replace_text_version_2(doc, 'RESIDENTREASON', reason.upper())
    replace_text_version_2(doc, 'RESIDENTSTAT', civil.lower())
    replace_text_version_2(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_2(doc, 'DAYSUFFIX', day_suffix)
    replace_text_version_2(doc, 'MONTHTODAY', month_name)
    replace_text_version_2(doc, 'YEARTODAY', str(datetime.now().year))
    replace_text_version_1(doc, 'KAPITANNAME', captain.upper())
    replace_text_version_1(doc, 'TAXCERTNUMBER', str(ctc))
    replace_text_version_1(doc, 'TAXCERTDATE', str(ctc_date))
    replace_text_version_2(doc, 'DATEISSUED', post_date_today)

    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Certificate of Good Moral Character")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Certificate of Good Moral Character")
    document_name = f"{fullName} - Certificate of Good Moral Character.docx"
    pdf_name = f"{fullName} - Certificate of Good Moral Character.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    # Save the document
    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_jobSeeker(fullName1, fullName2, purok, age, gender, witness, captain):
    doc_path = resources_path('assets\\templates\\Barangay_Certification.docx')  # Adjusted path
    doc = Document(doc_path)

    personal_pronoun = ""
    personal_pronoun_2 = ""
    pronoun = ""

    if gender.lower() == 'male':
        personal_pronoun = "him"
        personal_pronoun_2 = "his"
        pronoun = "he"
    elif gender.lower() == 'female':
        personal_pronoun = "her"
        pronoun = "she"

    day_suffix = get_ordinal_suffix(str_dayToday)

    replace_text_version_1(doc, 'RESIDENTNAME', fullName1.upper())
    replace_text_version_1(doc, 'RESIDENTSECNAME', fullName2.upper())
    replace_text_version_2(doc, 'RESIDENTPUROK', purok.capitalize())
    replace_text_version_2(doc, 'RESIDENTAGE', str(age))
    replace_text_version_2(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_2(doc, 'DAYSUFFIX', day_suffix)
    replace_text_version_2(doc, 'MONTHTODAY', month_name)
    replace_text_version_2(doc, 'YEARTODAY', str(datetime.now().year))
    replace_text_version_1(doc, 'ONEYEARFROMTODAY', str_yearFromToday)
    replace_text_version_1(doc, 'KAPITANNAME', f"HON. {captain.upper()}")
    replace_text_version_1(doc, 'KAGAWADNAME', witness.upper())
    replace_text_version_1(doc, 'DATETODAY', f'{month_name} {str_dayToday}, {str_yearToday}')
    replace_text_version_2(doc, 'PERSONALPRONOUN', personal_pronoun)
    replace_text_version_1(doc, 'PERSONALZPRONOUN', personal_pronoun_2)
    replace_text_version_1(doc, 'PSRONOUN', pronoun)

    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/First Time Jobseekers")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/First Time Jobseekers")
    document_name = f"{fullName1} - First Time Jobseekers and Oath.docx"
    pdf_name = f"{fullName1} - First Time Jobseekers and Oath.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    # Save the document
    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_barangayClearance(fullName, purok, gender, reason, ctc, ctc_date):
    doc_path = resources_path('assets\\templates\\Barangay_Clearance.docx')  # Adjusted path
    doc = Document(doc_path)

    personal_pronoun = ""
    pronoun = ""

    if gender.lower() == 'male':
        personal_pronoun = "him"
        pronoun = "he"
    elif gender.lower() == 'female':
        personal_pronoun = "her"
        pronoun = "she"

    day_suffix = get_ordinal_suffix(str_dayToday)
    kagawad_list = get_barangay_kagawad()

    replace_text_version_1(doc, 'KAPITANNAME', get_barangay_captain("name").upper())
    replace_text_version_1(doc, 'BKGWD1', kagawad_list[0].upper())
    replace_text_version_1(doc, 'BKGWD2', kagawad_list[1].upper())
    replace_text_version_1(doc, 'BKGWD3', kagawad_list[2].upper())
    replace_text_version_1(doc, 'BKGWD4', kagawad_list[3].upper())
    replace_text_version_1(doc, 'BKGWD5', kagawad_list[4].upper())
    replace_text_version_1(doc, 'BKGWD6', kagawad_list[5].upper())
    replace_text_version_1(doc, 'BKGWD7', kagawad_list[6].upper())
    replace_text_version_1(doc, 'BSCRTRY', get_barangay_secretary().upper())
    replace_text_version_1(doc, 'BTRSRR', get_barangay_treasurer().upper())
    replace_text_version_2(doc, 'RESIDENTNAME', fullName.upper())
    replace_text_version_2(doc, 'RESIDENTPUROK', purok.capitalize())
    replace_text_version_2(doc, 'RESIDENTREASON', reason.upper())
    replace_text_version_2(doc, 'PSRONOUN', pronoun)
    replace_text_version_2(doc, 'PERSONALPRONOUN', personal_pronoun)
    replace_text_version_2(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_2(doc, 'DAYSUFFIX', day_suffix)
    replace_text_version_1(doc, 'MONTHTODAY', month_name)
    replace_text_version_1(doc, 'YEARTODAY', str_yearToday)
    replace_text_version_2(doc, 'CTCNUMBER', str(ctc))
    replace_text_version_1(doc, 'TAXCERTDATE', str(ctc_date))
    replace_text_version_2(doc, 'DATETODAY', post_date_today)

    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Barangay Clearance")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Barangay Clearance")
    document_name = f"{fullName} - Barangay Clearance.docx"
    pdf_name = f"{fullName} - Barangay Clearance.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    # Save the document
    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_barangayClearance_on_duty(fullName, purok, gender, reason, ctc, ctc_date, officer):
    doc_path = resources_path('assets\\templates\\Barangay_Clearance_on_duty.docx')  # Adjusted path
    doc = Document(doc_path)

    personal_pronoun = ""
    pronoun = ""

    if gender.lower() == 'male':
        personal_pronoun = "him"
        pronoun = "he"
    elif gender.lower() == 'female':
        personal_pronoun = "her"
        pronoun = "she"

    day_suffix = get_ordinal_suffix(str_dayToday)
    kagawad_list = get_barangay_kagawad()

    replace_text_version_1(doc, 'KAPITANNAME', get_barangay_captain("name").upper())
    replace_text_version_1(doc, 'BKGWD1', kagawad_list[0].upper())
    replace_text_version_1(doc, 'BKGWD2', kagawad_list[1].upper())
    replace_text_version_1(doc, 'BKGWD3', kagawad_list[2].upper())
    replace_text_version_1(doc, 'BKGWD4', kagawad_list[3].upper())
    replace_text_version_1(doc, 'BKGWD5', kagawad_list[4].upper())
    replace_text_version_1(doc, 'BKGWD6', kagawad_list[5].upper())
    replace_text_version_1(doc, 'BKGWD7', kagawad_list[6].upper())
    replace_text_version_1(doc, 'BSCRTRY', get_barangay_secretary().upper())
    replace_text_version_1(doc, 'BTRSRR', get_barangay_treasurer().upper())
    replace_text_version_1(doc, 'BKGWDNAME', officer.upper())
    replace_text_version_2(doc, 'RESIDENTNAME', fullName.upper())
    replace_text_version_2(doc, 'RESIDENTPUROK', purok.capitalize())
    replace_text_version_2(doc, 'RESIDENTREASON', reason.upper())
    replace_text_version_2(doc, 'PSRONOUN', pronoun)
    replace_text_version_2(doc, 'PERSONALPRONOUN', personal_pronoun)
    replace_text_version_2(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_2(doc, 'DAYSUFFIX', day_suffix)
    replace_text_version_1(doc, 'MONTHTODAY', month_name)
    replace_text_version_1(doc, 'YEARTODAY', str_yearToday)
    replace_text_version_2(doc, 'CTCNUMBER', str(ctc))
    replace_text_version_1(doc, 'TAXCERTDATE', str(ctc_date))
    replace_text_version_2(doc, 'DATETODAY', post_date_today)

    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Barangay Clearance")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Barangay Clearance")
    document_name = f"{fullName} - Barangay Clearance.docx"
    pdf_name = f"{fullName} - Barangay Clearance.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    # Save the document
    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_barangayResidency(fullName, purok, gender, reason):
    doc_path = resources_path('assets\\templates\\Barangay_Residency.docx')  # Adjusted path
    doc = Document(doc_path)

    personal_pronoun = ""
    pronoun = ""

    if gender.lower() == 'male':
        personal_pronoun = "him"
        pronoun = "he"
    elif gender.lower() == 'female':
        personal_pronoun = "her"
        pronoun = "she"

    day_suffix = get_ordinal_suffix(str_dayToday)
    kagawad_list = get_barangay_kagawad()

    replace_text_version_1(doc, 'KAPITANNAME', get_barangay_captain('name').upper())
    replace_text_version_1(doc, 'BKGWD1', kagawad_list[0].upper())
    replace_text_version_1(doc, 'BKGWD2', kagawad_list[1].upper())
    replace_text_version_1(doc, 'BKGWD3', kagawad_list[2].upper())
    replace_text_version_1(doc, 'BKGWD4', kagawad_list[3].upper())
    replace_text_version_1(doc, 'BKGWD5', kagawad_list[4].upper())
    replace_text_version_1(doc, 'BKGWD6', kagawad_list[5].upper())
    replace_text_version_1(doc, 'BKGWD7', kagawad_list[6].upper())
    replace_text_version_1(doc, 'BSCRTRY', get_barangay_secretary().upper())
    replace_text_version_1(doc, 'BTRSRR', get_barangay_treasurer().upper())
    replace_text_version_2(doc, 'RESIDENTNAME', fullName.upper())
    replace_text_version_2(doc, 'RESIDENTPUROK', purok.capitalize())
    replace_text_version_2(doc, 'RESIDENTREASON', reason.upper())
    replace_text_version_2(doc, 'PSRONOUN', pronoun)
    replace_text_version_2(doc, 'PERSONALPRONOUN', personal_pronoun)
    replace_text_version_2(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_2(doc, 'DAYSUFFIX', day_suffix)
    replace_text_version_1(doc, 'MONTHTODAY', month_name)
    replace_text_version_1(doc, 'YEARTODAY', str_yearToday)

    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Barangay Residency")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Barangay Residency")
    document_name = f"{fullName} - Barangay Residency.docx"
    pdf_name = f"{fullName} - Barangay Residency.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    # Save the document
    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_barangayResidency_on_duty(fullName, purok, gender, reason, officer):
    doc_path = resources_path('assets\\templates\\Barangay_Residency_on_duty.docx')  # Adjusted path
    doc = Document(doc_path)

    personal_pronoun = ""
    pronoun = ""

    if gender.lower() == 'male':
        personal_pronoun = "him"
        pronoun = "he"
    elif gender.lower() == 'female':
        personal_pronoun = "her"
        pronoun = "she"

    day_suffix = get_ordinal_suffix(str_dayToday)
    kagawad_list = get_barangay_kagawad()

    replace_text_version_1(doc, 'KAPITANNAME', get_barangay_captain('name').upper())
    replace_text_version_1(doc, 'BKGWD1', kagawad_list[0].upper())
    replace_text_version_1(doc, 'BKGWD2', kagawad_list[1].upper())
    replace_text_version_1(doc, 'BKGWD3', kagawad_list[2].upper())
    replace_text_version_1(doc, 'BKGWD4', kagawad_list[3].upper())
    replace_text_version_1(doc, 'BKGWD5', kagawad_list[4].upper())
    replace_text_version_1(doc, 'BKGWD6', kagawad_list[5].upper())
    replace_text_version_1(doc, 'BKGWD7', kagawad_list[6].upper())
    replace_text_version_1(doc, 'BSCRTRY', get_barangay_secretary().upper())
    replace_text_version_1(doc, 'BTRSRR', get_barangay_treasurer().upper())
    replace_text_version_1(doc, 'BKGWDNAME', officer.upper())
    replace_text_version_2(doc, 'RESIDENTNAME', fullName.upper())
    replace_text_version_2(doc, 'RESIDENTPUROK', purok.capitalize())
    replace_text_version_2(doc, 'RESIDENTREASON', reason.upper())
    replace_text_version_2(doc, 'PSRONOUN', pronoun)
    replace_text_version_2(doc, 'PERSONALPRONOUN', personal_pronoun)
    replace_text_version_2(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_2(doc, 'DAYSUFFIX', day_suffix)
    replace_text_version_1(doc, 'MONTHTODAY', month_name)
    replace_text_version_1(doc, 'YEARTODAY', str_yearToday)

    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Barangay Residency")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Barangay Residency")
    document_name = f"{fullName} - Barangay Residency.docx"
    pdf_name = f"{fullName} - Barangay Residency.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    # Save the document
    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def check_document_exist(case_no, form):
    if form == "Form1":
        document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{case_no}")
        document_name = f"{case_no} - KP FORM No. 7 - COMPLAINT.pdf"
        document_path = os.path.join(document_save_path, document_name)
        absolute_document_path = os.path.abspath(document_path)
        if os.path.exists(absolute_document_path):
            return True
    if form == "Form2":
        document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{case_no}")
        document_name = f"{case_no} - KP FORM No. 9 - SUMMONS.pdf"
        document_path = os.path.join(document_save_path, document_name)
        absolute_document_path = os.path.abspath(document_path)
        if os.path.exists(absolute_document_path):
            return True
    if form == "Form4":
        document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{case_no}")
        document_name = f"{case_no} - KP FORM No. 12 - NOTICE OF HEARING.pdf"
        document_path = os.path.join(document_save_path, document_name)
        absolute_document_path = os.path.abspath(document_path)
        if os.path.exists(absolute_document_path):
            return True
    if form == "Form5":
        document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{case_no}")
        document_name = f"{case_no} - KP FORM No. 12 - FAILURE TO APPEAR.pdf"
        document_path = os.path.join(document_save_path, document_name)
        absolute_document_path = os.path.abspath(document_path)
        if os.path.exists(absolute_document_path):
            return True
    if form == "Form6":
        document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{case_no}")
        document_name = f"{case_no} - KP FORM No. 16 - AMICABLE SETTLEMENT.pdf"
        document_path = os.path.join(document_save_path, document_name)
        absolute_document_path = os.path.abspath(document_path)
        if os.path.exists(absolute_document_path):
            return True
    if form == "Form7":
        document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{case_no}")
        document_name = f"{case_no} - KP FORM No. 17 - REPUDIATION.pdf"
        document_path = os.path.join(document_save_path, document_name)
        absolute_document_path = os.path.abspath(document_path)
        if os.path.exists(absolute_document_path):
            return True
    if form == "Form8":
        document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{case_no}")
        document_name = f"{case_no} - KP FORM No. 13 - NOTICE OF HEARING.pdf"
        document_path = os.path.join(document_save_path, document_name)
        absolute_document_path = os.path.abspath(document_path)
        if os.path.exists(absolute_document_path):
            return True
    if form == "Form9":
        document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{case_no}")
        document_name = f"{case_no} - Certificate To File Action.pdf"
        document_path = os.path.join(document_save_path, document_name)
        absolute_document_path = os.path.abspath(document_path)
        if os.path.exists(absolute_document_path):
            return True
    if form == "Form10":
        document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{case_no}")
        document_name = f"{case_no} - Subpoena.pdf"
        document_path = os.path.join(document_save_path, document_name)
        absolute_document_path = os.path.abspath(document_path)
        if os.path.exists(absolute_document_path):
            return True
    if form == "Form11":
        document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{case_no}")
        document_name = f"{case_no} - NOTICE TO CHOSEN PANGKAT MEMBER.pdf"
        document_path = os.path.join(document_save_path, document_name)
        absolute_document_path = os.path.abspath(document_path)
        if os.path.exists(absolute_document_path):
            return True
    if form == "Form12":
        document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{case_no}")
        document_name = f"{case_no} - KP FORM No. 22 - BAR COUNTERCLAIM.pdf"
        document_path = os.path.join(document_save_path, document_name)
        absolute_document_path = os.path.abspath(document_path)
        if os.path.exists(absolute_document_path):
            return True
    if form == "Form13":
        document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{case_no}")
        document_name = f"{case_no} - KP FORM No. 21 - BAR ACTION.pdf"
        document_path = os.path.join(document_save_path, document_name)
        absolute_document_path = os.path.abspath(document_path)
        if os.path.exists(absolute_document_path):
            return True


def do_form_1(
        values,
        complainants,
        respondents,
        complainants_address,
        respondents_address,
        reason,
        relief
        ):
    
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - KP FORM No. 7 - COMPLAINT.docx"
    pdf_name = f"{values[2]} - KP FORM No. 7 - COMPLAINT.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    temp_date = datetime.strptime(values[1], '%d-%m-%Y')

    blotter_case_no = values[2]
    blotter_day = temp_date.strftime('%d')
    blotter_day = int(temp_date.strftime('%d'))
    blotter_month_number = temp_date.strftime('%m')
    blotter_month_word = temp_date.strftime('%B')
    blotter_year = temp_date.strftime('%Y')
    blotter_reason = values[-1]
    blotter_complainants = []
    for item in complainants:
        blotter_complainants.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_respondents = []
    for item in respondents:
        blotter_respondents.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_complainants_purok = complainants_address[2]
    blotter_complainants_barangay = complainants_address[3]
    blotter_complainants_city = complainants_address[4]
    blotter_complainants_province = complainants_address[5]
    blotter_respondents_purok = respondents_address[2]
    blotter_respondents_barangay = respondents_address[3]
    blotter_respondents_city = respondents_address[4]
    blotter_respondents_province = respondents_address[5]

    barangay_captain = str(get_barangay_captain("name")).upper()

    blotter_day_suffix = get_ordinal_suffix(int(blotter_day))

    doc_path = resources_path('assets/templates/KP_Form_No_7_Complainant.docx')  # Adjusted path
    doc = Document(doc_path)

    if len(blotter_complainants) > 1:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()}")

    if len(blotter_respondents) > 1:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()}")

    replace_text_version_1(doc, "CASENUMB", blotter_case_no)
    replace_text_version_1(doc, "COMPREASON", blotter_reason)

    replace_text_version_1(doc, "COMPPUROK", blotter_complainants_purok)
    replace_text_version_1(doc, "COMPBARANGAY", blotter_complainants_barangay)
    replace_text_version_1(doc, "COMPCITY", blotter_complainants_city)
    replace_text_version_1(doc, "COMPPROVINCE", blotter_complainants_province)

    replace_text_version_1(doc, "RESPPUROK", blotter_respondents_purok)
    replace_text_version_1(doc, "RESPBARANGAY", blotter_respondents_barangay)
    replace_text_version_1(doc, "RESPCITY", blotter_respondents_city)
    replace_text_version_1(doc, "RESPPROVINCE", blotter_respondents_province)

    replace_text_version_1(doc, "COMPELABORATE", reason)
    replace_text_version_1(doc, "COMPRELIEF", relief)

    replace_text_version_1(doc, "DAYFILED", str(blotter_day))
    replace_text_version_1(doc, "DAYFILLEDSUFFIX", blotter_day_suffix)
    replace_text_version_1(doc, "MONTHFILED", blotter_month_word)
    replace_text_version_1(doc, "YEARFILED", blotter_year)

    replace_text_version_1(doc, 'KAPITANNAME', barangay_captain)

    # Save the document
    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_form_2(
    values,
    complainants,
    respondents,
    complainants_address,
    respondents_address,
    summon_date,
    summon_time,
    summon_time_period
):
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - KP FORM No. 9 - SUMMONS.docx"
    pdf_name = f"{values[2]} - KP FORM No. 9 - SUMMONS.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    temp_date = datetime.strptime(values[1], '%d-%m-%Y')

    blotter_case_no = values[2]
    blotter_day = temp_date.strftime('%d')
    blotter_month_number = temp_date.strftime('%m')
    blotter_month_word = temp_date.strftime('%B')
    blotter_year = temp_date.strftime('%Y')
    blotter_reason = values[-1]
    blotter_complainants = []
    for item in complainants:
        blotter_complainants.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_respondents = []
    for item in respondents:
        blotter_respondents.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_complainants_purok = complainants_address[2]
    blotter_complainants_barangay = complainants_address[3]
    blotter_complainants_city = complainants_address[4]
    blotter_complainants_province = complainants_address[5]
    blotter_respondents_purok = respondents_address[2]
    blotter_respondents_barangay = respondents_address[3]
    blotter_respondents_city = respondents_address[4]
    blotter_respondents_province = respondents_address[5]

    temp_summon_date = datetime.strptime(str(summon_date), '%Y-%m-%d')
    summon_day = int(temp_summon_date.strftime('%d'))
    summon_month_number = temp_summon_date.strftime('%m')
    summon_month_word = temp_summon_date.strftime('%B')
    summon_year = temp_summon_date.strftime('%Y')

    barangay_captain = str(get_barangay_captain("name")).upper()

    summon_day_suffix = get_ordinal_suffix(int(summon_day))

    day_suffix = get_ordinal_suffix(str_dayToday)

    doc_path = resources_path('assets/templates/KP_Form_No_9_Summons.docx')  # Adjusted path
    doc = Document(doc_path)

    time_period = ""
    if summon_time_period.upper() == "AM":
        time_period = "morning"
    else:
        time_period = "afternoon"

    if len(blotter_complainants) > 1:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()}")

    if len(blotter_respondents) > 1:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()}")

    replace_text_version_1(doc, "CASENUMB", str(blotter_case_no))
    replace_text_version_1(doc, "COMPREASON", str(blotter_reason))

    replace_text_version_1(doc, "COMPPUROK", str(blotter_complainants_purok))
    replace_text_version_1(doc, "COMPBARANGAY", str(blotter_complainants_barangay))
    replace_text_version_1(doc, "COMPCITY", str(blotter_complainants_city))
    replace_text_version_1(doc, "COMPPROVINCE", str(blotter_complainants_province))

    replace_text_version_1(doc, "RESPPUROK", str(blotter_respondents_purok))
    replace_text_version_1(doc, "RESPBARANGAY", str(blotter_respondents_barangay))
    replace_text_version_1(doc, "RESPCITY", str(blotter_respondents_city))
    replace_text_version_1(doc, "RESPPROVINCE", str(blotter_respondents_province))

    replace_text_version_2(doc, 'SUMMONDAY', str(summon_day))
    replace_text_version_2(doc, 'SUMMONSUFFIX', str(summon_day_suffix))
    replace_text_version_1(doc, 'SUMMONMONTH', str(summon_month_word))
    replace_text_version_1(doc, 'SUMMMONYEAR', str(summon_year))
    
    replace_text_version_1(doc, 'SUMMONTIME', str(summon_time))
    replace_text_version_1(doc, 'SUMMONPERIODS', str(time_period))

    replace_text_version_2(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_2(doc, 'DAYSUFFIX', str(day_suffix))
    replace_text_version_1(doc, 'MONTHTODAY', str(month_name))
    replace_text_version_1(doc, 'YEARTODAY', str(str_yearToday))
    
    replace_text_version_1(doc, 'KAPITANNAME', str(barangay_captain))

    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_form_3(
    values,
    complainants,
    respondents,
    officer
):
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - KP FORM No. 9 - OFFICERS RETURN.docx"
    pdf_name = f"{values[2]} - KP FORM No. 9 - OFFICERS RETURN.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    blotter_case_no = values[2]
    blotter_complainants_names = []
    blotter_respondents_names = []

    barangay_captain = str(get_barangay_captain("name")).upper()

    day_suffix = get_ordinal_suffix(str_dayToday)

    for item in complainants:
        blotter_complainants_names.append(
            get_resident_name(item[2], item[3], item[4], item[5])
        )
    for item in respondents:
        blotter_respondents_names.append(
            get_resident_name(item[2], item[3], item[4], item[5])
        )

    if len(blotter_complainants_names) == 1:
        blotter_complainants_names.append("                            ")
    if len(blotter_respondents_names) == 1:
        blotter_respondents_names.append("                            ")
    
    doc_path = resources_path('assets/templates/KP_Form_No_9_Officers_Return.docx')  # Adjusted path
    doc = Document(doc_path)

    replace_text_version_1(doc, "{{RESP1PERM}}", blotter_respondents_names[0].upper())
    replace_text_version_1(doc, "{{RESP2PERM}}", blotter_respondents_names[1].upper())
    replace_text_version_2(doc, "DAYTODAY", str(str_dayToday))
    replace_text_version_2(doc, "DAYSUFFIX", str(day_suffix))
    replace_text_version_2(doc, "MONTHTODAY", str(month_name))
    replace_text_version_2(doc, "YEARTODAY", str(str_yearToday))
    replace_text_version_1(doc, "OFFICERNAME", str(officer).upper())

    if len(blotter_respondents_names) > 2:
        add_extended_paragraphs(doc, "EXTENDEDSENT", blotter_respondents_names)
    else:
        replace_text_version_2(doc, "EXTENDEDSENT", "by:")

    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_form_4(
        values,
        complainants,
        respondents,
        complainants_address,
        hearing_date,
        hearing_time,
        hearing_period
    ):
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - KP FORM No. 12 - NOTICE OF HEARING.docx"
    pdf_name = f"{values[2]} - KP FORM No. 12 - NOTICE OF HEARING.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    complainants_purok = complainants_address[2]
    complainants_barangay = complainants_address[3]
    complainants_city = complainants_address[4]
    complainants_province = complainants_address[5]

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    blotter_case_no = values[2]
    blotter_respondents_names = []

    complainants_name = ""

    temp_hearing_date = datetime.strptime(str(hearing_date), '%Y-%m-%d')
    hearing_day = int(temp_hearing_date.strftime('%d'))
    hearing_month_number = temp_hearing_date.strftime('%m')
    hearing_month_word = temp_hearing_date.strftime('%B')
    hearing_year = temp_hearing_date.strftime('%Y')

    hearing_day_suffix = get_ordinal_suffix(int(hearing_day))
    day_suffix = get_ordinal_suffix(str_dayToday)

    barangay_captain = str(get_barangay_captain("name")).upper()

    time_period = ""
    if hearing_period.upper() == "AM":
        time_period = "morning"
    else:
        time_period = "afternoon"

    if len(complainants) > 1:
        complainants_name = f"{get_resident_name(
            complainants[0][2],
            complainants[0][3],
            complainants[0][4],
            complainants[0][5]
        ).upper()} et al."
    else:
        complainants_name = get_resident_name(
            complainants[0][2],
            complainants[0][3],
            complainants[0][4],
            complainants[0][5]
        ).upper()

    for item in respondents:
        blotter_respondents_names.append(
            get_resident_name(
                item[2],
                item[3],
                item[4],
                item[5]
            )
        )

    respondents_name = ", ".join(map(str, blotter_respondents_names))
    respondents_name = respondents_name.upper()
    
    doc_path = resources_path('assets/templates/KP_Form_No_12_Notice_Of_Hearing.docx')  # Adjusted path
    doc = Document(doc_path)

    replace_text_version_1(doc, "COMPNAME", str(complainants_name))
    replace_text_version_1(doc, "COMPPUROK", str(complainants_purok))
    replace_text_version_1(doc, "COMPBARANGAY", str(complainants_barangay))
    replace_text_version_1(doc, "COMPCITY", str(complainants_city))
    replace_text_version_1(doc, "COMPPROVINCE", str(complainants_province))
    replace_text_version_1(doc, "HEARDAY", str(hearing_day))
    replace_text_version_1(doc, "HEARSUFFIX", str(hearing_day_suffix))
    replace_text_version_1(doc, "HEARMONTH", str(hearing_month_word))
    replace_text_version_1(doc, "HEARYEAR", str(hearing_year))
    replace_text_version_1(doc, "HEARTIME", str(hearing_time))
    replace_text_version_1(doc, "SUMMONPERIODS", str(time_period))
    replace_text_version_1(doc, "RESPNAME", str(respondents_name))
    replace_text_version_2(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_2(doc, 'DAYSUFFIX', str(day_suffix))
    replace_text_version_1(doc, 'MONTHTODAY', str(month_name))
    replace_text_version_1(doc, 'YEARTODAY', str(str_yearToday))
    replace_text_version_1(doc, 'KAPITANNAME', str(barangay_captain))

    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_form_5(
    values,
    respondents,
    respondents_address,
    hearing_date,
    hearing_time,
    hearing_period,
    hearing_previous_date,
    person_in_charge,
    person_role,
    complainants
):
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - KP FORM No. 12 - FAILURE TO APPEAR.docx"
    pdf_name = f"{values[2]} - KP FORM No. 12 - FAILURE TO APPEAR.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    blotter_case_no = values[2]
    blotter_respondents_name = ""
    blotter_respondents_purok = respondents_address[2]
    blotter_respondents_barangay = respondents_address[3]
    blotter_respondents_city = respondents_address[4]
    blotter_respondents_province = respondents_address[5]
    temp_hearing_date = datetime.strptime(str(hearing_date), '%Y-%m-%d')
    blotter_hearing_day = int(temp_hearing_date.strftime('%d'))
    blotter_hearing_day_suffix = get_ordinal_suffix(int(blotter_hearing_day))
    blotter_hearing_month_number = temp_hearing_date.strftime('%m')
    blotter_hearing_month_word = temp_hearing_date.strftime('%B')
    blotter_hearing_year = temp_hearing_date.strftime('%Y')
    blotter_hearing_period = ""
    blotter_person_in_charge = person_in_charge.upper()
    temp_hearing_previous_date = datetime.strptime(str(hearing_previous_date), '%Y-%m-%d')
    blotter_previous_hearing_date = f"{temp_hearing_previous_date.strftime('%B')} {temp_hearing_previous_date.strftime('%d')}, {temp_hearing_previous_date.strftime('%Y')}"
    day_suffix = get_ordinal_suffix(int(str_dayToday))
    table_complainants_names = []
    table_respondents_names = []

    for i in range(len(respondents)):
        table_respondents_names.append(
            get_resident_name(
                respondents[i][2],
                respondents[i][3],
                respondents[i][4],
                respondents[i][5]
            )
        )

    for i in range(len(complainants)):
        table_complainants_names.append(
            get_resident_name(
                complainants[i][2],
                complainants[i][3],
                complainants[i][4],
                complainants[i][5]
            )
        )

    if len(respondents) > 1:
        blotter_respondents_name = f"{get_resident_name(
            respondents[0][2],
            respondents[0][3],
            respondents[0][4],
            respondents[0][5]
        ).upper()} et al."
    else:
        blotter_respondents_name = f"{get_resident_name(
            respondents[0][2],
            respondents[0][3],
            respondents[0][4],
            respondents[0][5]
        ).upper()}"

    if hearing_period.upper() == "AM":
        blotter_hearing_period = "morning"
    else:
        blotter_hearing_period = "afternoon"

    doc_path = resources_path('assets/templates/KP_Form_No_12_Notice_Of_Hearing_Re_Failure.docx')
    doc = Document(doc_path)

    replace_text_version_1(doc, "REPNAME", str(blotter_respondents_name))
    replace_text_version_1(doc, "REPPUROK", str(blotter_respondents_purok))
    replace_text_version_1(doc, "REPPUROK", str(blotter_respondents_purok))
    replace_text_version_1(doc, "REPBARANGAY", str(blotter_respondents_barangay))
    replace_text_version_1(doc, "REPCITY", str(blotter_respondents_city))
    replace_text_version_1(doc, "REPPROVINCE", str(blotter_respondents_province))
    replace_text_version_1(doc, "HEARDAY", str(blotter_hearing_day))
    replace_text_version_1(doc, "HEARSUFFIX", str(blotter_hearing_day_suffix))
    replace_text_version_1(doc, "HEARMONTH", str(blotter_hearing_month_word))
    replace_text_version_1(doc, "HEARYEAR", str(blotter_hearing_year))
    replace_text_version_1(doc, "HEARTIME", str(hearing_time))
    replace_text_version_1(doc, "SUMMONPERIODS", str(blotter_hearing_period))
    replace_text_version_1(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_1(doc, 'DAYSUFFIX', str(day_suffix))
    replace_text_version_1(doc, 'MONTHTODAY', str(month_name))
    replace_text_version_1(doc, 'YEARTODAY', str(str_yearToday))
    replace_text_version_1(doc, 'PERINCHARGE', str(blotter_person_in_charge))
    replace_text_version_1(doc, 'PERROLE', str(person_role))
    replace_text_version_1(doc, 'HEARPREVDATE', str(blotter_previous_hearing_date))
    set_tables(doc, table_complainants_names, table_respondents_names, 1)

    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)

def do_form_6(
    values,
    complainants,
    respondents,
    complainants_address,
    respondents_address,
    settlement_statement
):
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - KP FORM No. 16 - AMICABLE SETTLEMENT.docx"
    pdf_name = f"{values[2]} - KP FORM No. 16 - AMICABLE SETTLEMENT.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    temp_date = datetime.strptime(values[1], '%d-%m-%Y')

    blotter_case_no = values[2]
    blotter_day = temp_date.strftime('%d')
    blotter_month_number = temp_date.strftime('%m')
    blotter_month_word = temp_date.strftime('%B')
    blotter_year = temp_date.strftime('%Y')
    blotter_reason = values[-1]
    blotter_complainants = []
    for item in complainants:
        blotter_complainants.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_respondents = []
    for item in respondents:
        blotter_respondents.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_complainants_purok = complainants_address[2]
    blotter_complainants_barangay = complainants_address[3]
    blotter_complainants_city = complainants_address[4]
    blotter_complainants_province = complainants_address[5]
    blotter_respondents_purok = respondents_address[2]
    blotter_respondents_barangay = respondents_address[3]
    blotter_respondents_city = respondents_address[4]
    blotter_respondents_province = respondents_address[5]

    day_suffix = get_ordinal_suffix(int(str_dayToday))

    table_complainants_names = []
    table_respondents_names = []

    for i in range(len(respondents)):
        table_respondents_names.append(
            get_resident_name(
                respondents[i][2],
                respondents[i][3],
                respondents[i][4],
                respondents[i][5]
            )
        )

    for i in range(len(complainants)):
        table_complainants_names.append(
            get_resident_name(
                complainants[i][2],
                complainants[i][3],
                complainants[i][4],
                complainants[i][5]
            )
        )

    barangay_captain = str(get_barangay_captain("name")).upper()

    blotter_day_suffix = get_ordinal_suffix(int(blotter_day))

    doc_path = resources_path('assets/templates/KP_Form_No_16_Amicable_Settlement.docx')  # Adjusted path
    doc = Document(doc_path)

    if len(blotter_complainants) > 1:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()}")

    if len(blotter_respondents) > 1:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()}")

    replace_text_version_1(doc, "CASENUMB", blotter_case_no)
    replace_text_version_1(doc, "COMPREASON", blotter_reason)

    replace_text_version_1(doc, "COMPPUROK", blotter_complainants_purok)
    replace_text_version_1(doc, "COMPBARANGAY", blotter_complainants_barangay)
    replace_text_version_1(doc, "COMPCITY", blotter_complainants_city)
    replace_text_version_1(doc, "COMPPROVINCE", blotter_complainants_province)

    replace_text_version_1(doc, "SETTSTATE", settlement_statement)

    replace_text_version_1(doc, "RESPPUROK", blotter_respondents_purok)
    replace_text_version_1(doc, "RESPBARANGAY", blotter_respondents_barangay)
    replace_text_version_1(doc, "RESPCITY", blotter_respondents_city)
    replace_text_version_1(doc, "RESPPROVINCE", blotter_respondents_province)

    replace_text_version_1(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_1(doc, 'DAYSUFFIX', str(day_suffix))
    replace_text_version_1(doc, 'MONTHTODAY', str(month_name))
    replace_text_version_1(doc, 'YEARTODAY', str(str_yearToday))

    replace_text_version_1(doc, 'KAPITANNAME', barangay_captain)

    set_tables(doc, table_complainants_names, table_respondents_names, 2)

    # Save the document
    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)
    

def do_form_7(
        values,
        complainants,
        respondents,
        complainants_address,
        respondents_address,
        fraud,
        violence,
        intimidation,
        in_charge_name,
        in_charge_role
        ):
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - KP FORM No. 17 - REPUDIATION.docx"
    pdf_name = f"{values[2]} - KP FORM No. 17 - REPUDIATION.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    temp_date = datetime.strptime(values[1], '%d-%m-%Y')

    blotter_case_no = values[2]
    blotter_day = temp_date.strftime('%d')
    blotter_month_number = temp_date.strftime('%m')
    blotter_month_word = temp_date.strftime('%B')
    blotter_year = temp_date.strftime('%Y')
    blotter_reason = values[-1]
    blotter_complainants = []
    for item in complainants:
        blotter_complainants.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_respondents = []
    for item in respondents:
        blotter_respondents.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_complainants_purok = complainants_address[2]
    blotter_complainants_barangay = complainants_address[3]
    blotter_complainants_city = complainants_address[4]
    blotter_complainants_province = complainants_address[5]
    blotter_respondents_purok = respondents_address[2]
    blotter_respondents_barangay = respondents_address[3]
    blotter_respondents_city = respondents_address[4]
    blotter_respondents_province = respondents_address[5]

    day_suffix = get_ordinal_suffix(int(str_dayToday))

    table_complainants_names = []
    table_respondents_names = []

    for i in range(len(respondents)):
        table_respondents_names.append(
            get_resident_name(
                respondents[i][2],
                respondents[i][3],
                respondents[i][4],
                respondents[i][5]
            )
        )

    for i in range(len(complainants)):
        table_complainants_names.append(
            get_resident_name(
                complainants[i][2],
                complainants[i][3],
                complainants[i][4],
                complainants[i][5]
            )
        )

    barangay_captain = str(get_barangay_captain("name")).upper()

    blotter_day_suffix = get_ordinal_suffix(int(blotter_day))

    doc_path = resources_path('assets/templates/KP_Form_No_17_Repudiation.docx')  # Adjusted path
    doc = Document(doc_path)

    if len(blotter_complainants) > 1:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()}")

    if len(blotter_respondents) > 1:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()}")

    replace_text_version_1(doc, "CASENUMB", blotter_case_no)
    replace_text_version_1(doc, "COMPREASON", blotter_reason)

    replace_text_version_1(doc, "COMPPUROK", blotter_complainants_purok)
    replace_text_version_1(doc, "COMPBARANGAY", blotter_complainants_barangay)
    replace_text_version_1(doc, "COMPCITY", blotter_complainants_city)
    replace_text_version_1(doc, "COMPPROVINCE", blotter_complainants_province)

    replace_text_version_1(doc, "RESPPUROK", blotter_respondents_purok)
    replace_text_version_1(doc, "RESPBARANGAY", blotter_respondents_barangay)
    replace_text_version_1(doc, "RESPCITY", blotter_respondents_city)
    replace_text_version_1(doc, "RESPPROVINCE", blotter_respondents_province)

    replace_text_version_1(doc, "STATEFRAUD", fraud)
    replace_text_version_1(doc, "STATEVIOLENCE", violence)
    replace_text_version_1(doc, "STATEINTIMID", intimidation)

    replace_text_version_1(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_1(doc, 'DAYSUFFIX', str(day_suffix))
    replace_text_version_1(doc, 'MONTHTODAY', str(month_name))
    replace_text_version_1(doc, 'YEARTODAY', str(str_yearToday))

    replace_text_version_1(doc, 'PEPCHARGE', in_charge_name.upper())
    replace_text_version_1(doc, 'PEPROLE', in_charge_role)

    replace_text_version_1(doc, 'KAPITANNAME', barangay_captain)

    set_tables(doc, table_complainants_names, table_respondents_names, 3)

    # Save the document
    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)
    

def do_form_8(
    values,
    complainants,
    respondents,
    complainants_address,
    respondents_address,
    summon_date,
    summon_time,
    summon_time_period
):
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - KP FORM No. 13 - NOTICE OF HEARING.docx"
    pdf_name = f"{values[2]} - KP FORM No. 13 - NOTICE OF HEARING.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    temp_date = datetime.strptime(values[1], '%d-%m-%Y')

    blotter_case_no = values[2]
    blotter_day = temp_date.strftime('%d')
    blotter_month_number = temp_date.strftime('%m')
    blotter_month_word = temp_date.strftime('%B')
    blotter_year = temp_date.strftime('%Y')
    blotter_reason = values[-1]
    blotter_complainants = []
    for item in complainants:
        blotter_complainants.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_respondents = []
    for item in respondents:
        blotter_respondents.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_complainants_purok = complainants_address[2]
    blotter_complainants_barangay = complainants_address[3]
    blotter_complainants_city = complainants_address[4]
    blotter_complainants_province = complainants_address[5]
    blotter_respondents_purok = respondents_address[2]
    blotter_respondents_barangay = respondents_address[3]
    blotter_respondents_city = respondents_address[4]
    blotter_respondents_province = respondents_address[5]

    temp_summon_date = datetime.strptime(str(summon_date), '%Y-%m-%d')
    summon_day = int(temp_summon_date.strftime('%d'))
    summon_month_number = temp_summon_date.strftime('%m')
    summon_month_word = temp_summon_date.strftime('%B')
    summon_year = temp_summon_date.strftime('%Y')

    barangay_captain = str(get_barangay_captain("name")).upper()

    summon_day_suffix = get_ordinal_suffix(int(summon_day))

    day_suffix = get_ordinal_suffix(str_dayToday)

    doc_path = resources_path('assets/templates/KP_Form_No_13_Notice_Of_Hearing.docx')  # Adjusted path
    doc = Document(doc_path)

    time_period = ""
    if summon_time_period.upper() == "AM":
        time_period = "morning"
    else:
        time_period = "afternoon"

    if len(blotter_complainants) > 1:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()}")

    if len(blotter_respondents) > 1:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()}")

    replace_text_version_1(doc, "CASENUMB", str(blotter_case_no))
    replace_text_version_1(doc, "COMPREASON", str(blotter_reason))

    replace_text_version_1(doc, "COMPPUROK", str(blotter_complainants_purok))
    replace_text_version_1(doc, "COMPBARANGAY", str(blotter_complainants_barangay))
    replace_text_version_1(doc, "COMPCITY", str(blotter_complainants_city))
    replace_text_version_1(doc, "COMPPROVINCE", str(blotter_complainants_province))

    replace_text_version_1(doc, "RESPPUROK", str(blotter_respondents_purok))
    replace_text_version_1(doc, "RESPBARANGAY", str(blotter_respondents_barangay))
    replace_text_version_1(doc, "RESPCITY", str(blotter_respondents_city))
    replace_text_version_1(doc, "RESPPROVINCE", str(blotter_respondents_province))

    replace_text_version_2(doc, 'SUMMONDAY', str(summon_day))
    replace_text_version_2(doc, 'SUMMONSUFFIX', str(summon_day_suffix))
    replace_text_version_1(doc, 'SUMMONMONTH', str(summon_month_word))
    replace_text_version_1(doc, 'SUMMMONYEAR', str(summon_year))
    
    replace_text_version_1(doc, 'SUMMONTIME', str(summon_time))
    replace_text_version_1(doc, 'SUMMONPERIOD', str(time_period))

    replace_text_version_2(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_2(doc, 'DAYSUFFIX', str(day_suffix))
    replace_text_version_1(doc, 'MONTHTODAY', str(month_name))
    replace_text_version_1(doc, 'YEARTODAY', str(str_yearToday))
    
    replace_text_version_1(doc, 'KAPITANNAME', str(barangay_captain))

    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_form_9(
    values,
    complainants,
    respondents,
    complainants_address,
    respondents_address,
    member,
    secretary,
    chairman
):
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - Certificate To File Action.docx"
    pdf_name = f"{values[2]} - Certificate To File Action.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    temp_date = datetime.strptime(values[1], '%d-%m-%Y')

    blotter_case_no = values[2]
    blotter_reason = values[-1]
    blotter_complainants = []
    for item in complainants:
        blotter_complainants.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_respondents = []
    for item in respondents:
        blotter_respondents.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_complainants_purok = complainants_address[2]
    blotter_complainants_barangay = complainants_address[3]
    blotter_complainants_city = complainants_address[4]
    blotter_complainants_province = complainants_address[5]
    blotter_respondents_purok = respondents_address[2]
    blotter_respondents_barangay = respondents_address[3]
    blotter_respondents_city = respondents_address[4]
    blotter_respondents_province = respondents_address[5]

    day_suffix = get_ordinal_suffix(str_dayToday)

    doc_path = resources_path('assets/templates/Certificate_To_File_Action.docx')  # Adjusted path
    doc = Document(doc_path)


    if len(blotter_complainants) > 1:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()}")

    if len(blotter_respondents) > 1:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()}")

    replace_text_version_1(doc, "COMPPUROK", str(blotter_complainants_purok))
    replace_text_version_1(doc, "COMPBARANGAY", str(blotter_complainants_barangay))
    replace_text_version_1(doc, "COMPCITY", str(blotter_complainants_city))
    replace_text_version_1(doc, "COMPPROVINCE", str(blotter_complainants_province))

    replace_text_version_1(doc, "CASENUMB", str(blotter_case_no))
    replace_text_version_1(doc, "COMPREASON", str(blotter_reason))

    replace_text_version_1(doc, "RESPPUROK", str(blotter_respondents_purok))
    replace_text_version_1(doc, "RESPBARANGAY", str(blotter_respondents_barangay))
    replace_text_version_1(doc, "RESPCITY", str(blotter_respondents_city))
    replace_text_version_1(doc, "RESPPROVINCE", str(blotter_respondents_province))

    replace_text_version_1(doc, "PKTMEM", str(member).upper())
    replace_text_version_1(doc, "PKTSEC", str(secretary).upper())
    replace_text_version_1(doc, "PKTCHAIR", str(chairman).upper())

    replace_text_version_2(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_2(doc, 'DAYSUFFIX', str(day_suffix))
    replace_text_version_1(doc, 'MONTHTODAY', str(month_name))
    replace_text_version_1(doc, 'YEARTODAY', str(str_yearToday))

    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_form_10(
    values,
    complainants,
    respondents,
    complainants_address,
    respondents_address,
    summon_name,
    summon_purok,
    summon_barangay,
    summon_city,
    summon_province,
    summon_date,
    summon_time,
    summon_time_period,
    in_charge_name,
    in_charge_role
):
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - Subpoena.docx"
    pdf_name = f"{values[2]} - Subpoena.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    temp_date = datetime.strptime(values[1], '%d-%m-%Y')

    blotter_case_no = values[2]
    blotter_day = temp_date.strftime('%d')
    blotter_month_number = temp_date.strftime('%m')
    blotter_month_word = temp_date.strftime('%B')
    blotter_year = temp_date.strftime('%Y')
    blotter_reason = values[-1]
    blotter_complainants = []
    for item in complainants:
        blotter_complainants.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_respondents = []
    for item in respondents:
        blotter_respondents.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_complainants_purok = complainants_address[2]
    blotter_complainants_barangay = complainants_address[3]
    blotter_complainants_city = complainants_address[4]
    blotter_complainants_province = complainants_address[5]
    blotter_respondents_purok = respondents_address[2]
    blotter_respondents_barangay = respondents_address[3]
    blotter_respondents_city = respondents_address[4]
    blotter_respondents_province = respondents_address[5]

    temp_summon_date = datetime.strptime(str(summon_date), '%Y-%m-%d')
    summon_day = int(temp_summon_date.strftime('%d'))
    summon_month_number = temp_summon_date.strftime('%m')
    summon_month_word = temp_summon_date.strftime('%B')
    summon_year = temp_summon_date.strftime('%Y')

    if in_charge_role == "Punong Barangay":
        in_charge_name = f"HON. {in_charge_name}"

    summon_day_suffix = get_ordinal_suffix(int(summon_day))

    day_suffix = get_ordinal_suffix(str_dayToday)

    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    doc_path = resources_path('assets/templates/Subpoena.docx')  # Adjusted path
    doc = Document(doc_path)

    time_period = ""
    if summon_time_period.upper() == "AM":
        time_period = "morning"
    else:
        time_period = "afternoon"

    if len(blotter_complainants) > 1:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()}")

    if len(blotter_respondents) > 1:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()}")

    replace_text_version_1(doc, "COMPPUROK", str(blotter_complainants_purok))
    replace_text_version_1(doc, "COMPBARANGAY", str(blotter_complainants_barangay))
    replace_text_version_1(doc, "COMPCITY", str(blotter_complainants_city))
    replace_text_version_1(doc, "COMPPROVINCE", str(blotter_complainants_province))

    replace_text_version_1(doc, "RESPPUROK", str(blotter_respondents_purok))
    replace_text_version_1(doc, "RESPBARANGAY", str(blotter_respondents_barangay))
    replace_text_version_1(doc, "RESPCITY", str(blotter_respondents_city))
    replace_text_version_1(doc, "RESPPROVINCE", str(blotter_respondents_province))

    replace_text_version_1(doc, "SUBNAME", str(summon_name).upper())
    replace_text_version_1(doc, "SUBPUROK", str(summon_purok))
    replace_text_version_1(doc, "SUBBARANGAY", str(summon_barangay))
    replace_text_version_1(doc, "SUBCITY", str(summon_city))
    replace_text_version_1(doc, "SUBPROVINCE", str(summon_province))

    replace_text_version_2(doc, 'SUMMONDAY', str(summon_day))
    replace_text_version_2(doc, 'SUMMONSUFFIX', str(summon_day_suffix))
    replace_text_version_1(doc, 'SUMMONMONTH', str(summon_month_word))
    replace_text_version_1(doc, 'SUMMMONYEAR', str(summon_year))
    
    replace_text_version_1(doc, 'SUMMONTIME', str(summon_time))
    replace_text_version_1(doc, 'SUMMONPERIOD', str(time_period))

    replace_text_version_1(doc, 'INCHARGENAME', in_charge_name.upper())
    replace_text_version_1(doc, 'INCHARGEROLE', in_charge_role)

    replace_text_version_2(doc, 'DAYTODAY', str(str_dayToday))
    replace_text_version_2(doc, 'DAYSUFFIX', str(day_suffix))
    replace_text_version_1(doc, 'MONTHTODAY', str(month_name))
    replace_text_version_1(doc, 'YEARTODAY', str(str_yearToday))

    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_form_11(
    values,
    complainants,
    respondents,
    complainants_address,
    respondents_address,
    summon_date,
    summon_time,
    summon_time_period,
    first_pangkat_member,
    second_pangkat_member,
    third_pangkat_member,
    first_pangkat_purok,
    second_pangkat_purok,
    third_pangkat_purok
):
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - NOTICE TO CHOSEN PANGKAT MEMBER.docx"
    pdf_name = f"{values[2]} - NOTICE TO CHOSEN PANGKAT MEMBER.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    temp_date = datetime.strptime(values[1], '%d-%m-%Y')

    blotter_case_no = values[2]
    blotter_reason = values[-1]
    blotter_complainants = []
    for item in complainants:
        blotter_complainants.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_respondents = []
    for item in respondents:
        blotter_respondents.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_complainants_purok = complainants_address[2]
    blotter_complainants_barangay = complainants_address[3]
    blotter_complainants_city = complainants_address[4]
    blotter_complainants_province = complainants_address[5]
    blotter_respondents_purok = respondents_address[2]
    blotter_respondents_barangay = respondents_address[3]
    blotter_respondents_city = respondents_address[4]
    blotter_respondents_province = respondents_address[5]

    temp_summon_date = datetime.strptime(str(summon_date), '%Y-%m-%d')
    summon_day = int(temp_summon_date.strftime('%d'))
    summon_month_word = temp_summon_date.strftime('%B')
    summon_year = temp_summon_date.strftime('%Y')

    barangay_captain = str(get_barangay_captain("name")).upper()

    doc_path = resources_path('assets/templates/Notice_To_Chosen_Pangkat.docx')  # Adjusted path
    doc = Document(doc_path)

    time_period = ""
    if summon_time_period.upper() == "AM":
        time_period = "morning"
    else:
        time_period = "afternoon"

    if len(blotter_complainants) > 1:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()}")

    if len(blotter_respondents) > 1:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()}")

    replace_text_version_1(doc, "CASENUMB", str(blotter_case_no))
    replace_text_version_1(doc, "COMPREASON", str(blotter_reason))

    replace_text_version_1(doc, "COMPPUROK", str(blotter_complainants_purok))
    replace_text_version_1(doc, "COMPBARANGAY", str(blotter_complainants_barangay))
    replace_text_version_1(doc, "COMPCITY", str(blotter_complainants_city))
    replace_text_version_1(doc, "COMPPROVINCE", str(blotter_complainants_province))

    replace_text_version_1(doc, "RESPPUROK", str(blotter_respondents_purok))
    replace_text_version_1(doc, "RESPBARANGAY", str(blotter_respondents_barangay))
    replace_text_version_1(doc, "RESPCITY", str(blotter_respondents_city))
    replace_text_version_1(doc, "RESPPROVINCE", str(blotter_respondents_province))

    replace_text_version_2(doc, 'DATETODAY', f"{str(month_name)} {str(str_dayToday)}, {str(str_yearToday)}")

    replace_text_version_2(doc, 'SESHDAY', str(summon_day))
    replace_text_version_2(doc, 'SESHMONTH', str(summon_month_word))
    replace_text_version_2(doc, 'SESHYEAR', str(summon_year))
    
    replace_text_version_2(doc, 'SESHTIME', str(summon_time))
    replace_text_version_2(doc, 'SESHPERIOD', str(time_period))
    
    replace_text_version_1(doc, 'FIRSTPKTMEM', str(first_pangkat_member))
    replace_text_version_1(doc, 'FIRSTPRK', str(first_pangkat_purok))

    replace_text_version_1(doc, 'SECONDPKTMEM', str(second_pangkat_member))
    replace_text_version_1(doc, '{{SECONDPRK}}', str(second_pangkat_purok))

    replace_text_version_1(doc, 'THIRDPKTMEM', str(third_pangkat_member))
    replace_text_version_1(doc, '{{THIRDPRK}}', str(third_pangkat_purok))

    replace_text_version_1(doc, 'KAPITANNAME', str(barangay_captain))

    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_form_12(
    values,
    complainants,
    respondents,
    complainants_address,
    respondents_address,
    secretary,
    chairman
):
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - KP FORM No. 22 - BAR COUNTERCLAIM.docx"
    pdf_name = f"{values[2]} - KP FORM No. 22 - BAR COUNTERCLAIM.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)
    
    temp_date = datetime.strptime(values[1], '%d-%m-%Y')

    blotter_case_no = values[2]
    blotter_day = temp_date.strftime('%d')
    blotter_month_number = temp_date.strftime('%m')
    blotter_month_word = temp_date.strftime('%B')
    blotter_year = temp_date.strftime('%Y')
    blotter_reason = values[-1]
    blotter_complainants = []
    for item in complainants:
        blotter_complainants.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_respondents = []
    for item in respondents:
        blotter_respondents.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_complainants_purok = complainants_address[2]
    blotter_complainants_barangay = complainants_address[3]
    blotter_complainants_city = complainants_address[4]
    blotter_complainants_province = complainants_address[5]
    blotter_respondents_purok = respondents_address[2]
    blotter_respondents_barangay = respondents_address[3]
    blotter_respondents_city = respondents_address[4]
    blotter_respondents_province = respondents_address[5]

    day_suffix = get_ordinal_suffix(str_dayToday)
    
    blotter_complainants_names = []
    blotter_respondents_names = []

    for item in complainants:
        blotter_complainants_names.append(
            get_resident_name(item[2], item[3], item[4], item[5])
        )
    for item in respondents:
        blotter_respondents_names.append(
            get_resident_name(item[2], item[3], item[4], item[5])
        )

    if len(blotter_complainants_names) == 1:
        blotter_complainants_names.append("                            ")
    if len(blotter_respondents_names) == 1:
        blotter_respondents_names.append("                            ")
    
    doc_path = resources_path('assets/templates/KP_Form_No_22_Bar_Counterclaim.docx')  # Adjusted path
    doc = Document(doc_path)

    if len(blotter_complainants) > 1:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()}")

    if len(blotter_respondents) > 1:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()}")

    replace_text_version_1(doc, "COMPPUROK", str(blotter_complainants_purok))
    replace_text_version_1(doc, "COMPBARANGAY", str(blotter_complainants_barangay))
    replace_text_version_1(doc, "COMPCITY", str(blotter_complainants_city))
    replace_text_version_1(doc, "COMPPROVINCE", str(blotter_complainants_province))

    replace_text_version_1(doc, "CASENUMB", str(blotter_case_no))
    replace_text_version_1(doc, "COMPREASON", str(blotter_reason))

    replace_text_version_1(doc, "RESPPUROK", str(blotter_respondents_purok))
    replace_text_version_1(doc, "RESPBARANGAY", str(blotter_respondents_barangay))
    replace_text_version_1(doc, "RESPCITY", str(blotter_respondents_city))
    replace_text_version_1(doc, "RESPPROVINCE", str(blotter_respondents_province))

    replace_text_version_1(doc, "{{RESP1PERM}}", blotter_respondents_names[0].upper())
    replace_text_version_1(doc, "{{RESP2PERM}}", blotter_respondents_names[1].upper())

    replace_text_version_1(doc, "PERSECRETARY", secretary.upper())
    replace_text_version_1(doc, "PERCHAIRMAN", chairman.upper())
    
    replace_text_version_2(doc, "DAYTODAY", str(str_dayToday))
    replace_text_version_2(doc, "DAYSUFFIX", str(day_suffix))
    replace_text_version_2(doc, "MONTHTODAY", str(month_name))
    replace_text_version_2(doc, "YEARTODAY", str(str_yearToday))

    if len(blotter_respondents_names) > 2:
        add_counterclaim_extended_paragraphs(doc, "EXTENDEDSENT", blotter_respondents_names)
    else:
        replace_text_version_2(doc, "EXTENDEDSENT", " have been found to have willfully failed or refused to appear without justifiable reason before the Punong Barangay/Pangkat ng Tagapagkasundo and therefore respondent/s is are barred from filing his/their counterclaim (if any) arising from the complaint in court/government office.")

    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)

def do_form_13(
    values,
    complainants,
    respondents,
    complainants_address,
    respondents_address,
    secretary,
    member,
    chairman
):
    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Blotter Documents/{values[2]}")
    document_name = f"{values[2]} - KP FORM No. 21 - BAR ACTION.docx"
    pdf_name = f"{values[2]} - KP FORM No. 21 - BAR ACTION.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)
    
    temp_date = datetime.strptime(values[1], '%d-%m-%Y')

    blotter_case_no = values[2]
    blotter_day = temp_date.strftime('%d')
    blotter_month_number = temp_date.strftime('%m')
    blotter_month_word = temp_date.strftime('%B')
    blotter_year = temp_date.strftime('%Y')
    blotter_reason = values[-1]
    blotter_complainants = []
    for item in complainants:
        blotter_complainants.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_respondents = []
    for item in respondents:
        blotter_respondents.append(get_resident_name(item[2], item[3], item[4], item[5]))
    blotter_complainants_purok = complainants_address[2]
    blotter_complainants_barangay = complainants_address[3]
    blotter_complainants_city = complainants_address[4]
    blotter_complainants_province = complainants_address[5]
    blotter_respondents_purok = respondents_address[2]
    blotter_respondents_barangay = respondents_address[3]
    blotter_respondents_city = respondents_address[4]
    blotter_respondents_province = respondents_address[5]

    day_suffix = get_ordinal_suffix(str_dayToday)
    
    blotter_complainants_names = []
    blotter_respondents_names = []

    for item in complainants:
        blotter_complainants_names.append(
            get_resident_name(item[2], item[3], item[4], item[5])
        )
    for item in respondents:
        blotter_respondents_names.append(
            get_resident_name(item[2], item[3], item[4], item[5])
        )
    
    doc_path = resources_path('assets/templates/KP_Form_No_21_Bar_Action.docx')  # Adjusted path
    doc = Document(doc_path)

    if len(blotter_complainants) > 1:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'COMPNAME', f"{blotter_complainants[0].upper()}")

    if len(blotter_respondents) > 1:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()} et al.")
    else:
        replace_text_version_1(doc, 'RESPNAME', f"{blotter_respondents[0].upper()}")

    replace_text_version_1(doc, "COMPPUROK", str(blotter_complainants_purok))
    replace_text_version_1(doc, "COMPBARANGAY", str(blotter_complainants_barangay))
    replace_text_version_1(doc, "COMPCITY", str(blotter_complainants_city))
    replace_text_version_1(doc, "COMPPROVINCE", str(blotter_complainants_province))

    replace_text_version_1(doc, "CASENUMB", str(blotter_case_no))
    replace_text_version_1(doc, "COMPREASON", str(blotter_reason))

    replace_text_version_1(doc, "RESPPUROK", str(blotter_respondents_purok))
    replace_text_version_1(doc, "RESPBARANGAY", str(blotter_respondents_barangay))
    replace_text_version_1(doc, "RESPCITY", str(blotter_respondents_city))
    replace_text_version_1(doc, "RESPPROVINCE", str(blotter_respondents_province))

    add_action_extended_paragraphs(doc, "EXTENDEDSENT", blotter_complainants_names, f"{blotter_month_word} {int(blotter_day)}, {blotter_year}")

    replace_text_version_1(doc, "PKTSEC", secretary.upper())
    replace_text_version_1(doc, "PKTMEM", member.upper())
    replace_text_version_1(doc, "PKTCHAIR", chairman.upper())
    
    replace_text_version_2(doc, "DAYTODAY", str(str_dayToday))
    replace_text_version_2(doc, "DAYSUFFIX", str(day_suffix))
    replace_text_version_2(doc, "MONTHTODAY", str(month_name))
    replace_text_version_2(doc, "YEARTODAY", str(str_yearToday))

    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def generate_reports(values, filter_values):
    doc_path = resources_path('assets/templates/residents_reports_copy.docx')  # Adjusted path
    doc = f_Document(doc_path)

    # PLACEHOLDER_TEXT = "{{TABLE_PLACEHOLDER}}"

    # for par in doc.paragraphs:
    #     if PLACEHOLDER_TEXT in par.text:
    #         table = replace_with_table(doc, par._p, par._parent, len(values), 4)
    #         for row in range(len(values)):
    #             table.cell(row, 0).text = values[row][0]
    #             table.cell(row, 1).text = str(values[row][1])
    #             table.cell(row, 2).text = str(values[row][2])
    #             table.cell(row, 3).text = str(values[row][3])
    #             table.cell(row, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #             table.cell(row, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #             table.cell(row, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #             table.cell(row, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #             table.cell(row, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #             table.cell(row, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #             table.cell(row, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #             table.cell(row, 3).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Reports")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Residents Profiling Documents/Reports")
    document_name = "Barangay Residents Reports.docx"
    pdf_name = "Barangay Residents Reports.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)

    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    # replace_text_version_1(doc, '{{REPORT_TITLE}}', title.upper())
    
    # add_reports_paragraphs(doc, "{{REPORT_TITLE}}", filter_values)

    set_report_filters(doc, filter_values)

    set_report_tables(doc, values, 2)

    replace_text_version_1(doc, 'BSCRTRY', get_barangay_secretary().upper())
    replace_text_version_1(doc, 'KAPITANNAME', get_barangay_captain("name").upper())

    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


def do_barangay_officials_list():
    doc_path = resources_path('assets\\templates\\barangay_officials.docx')  # Adjusted path
    doc = Document(doc_path)

    kagawad_list = get_barangay_kagawad()
    lupon_list = get_barangay_pangkat_member()


    replace_text_version_1(doc, 'KAPITANNAME', get_barangay_captain('name').upper())
    replace_text_version_1(doc, 'BKGWD1', kagawad_list[0].upper())
    replace_text_version_1(doc, 'BKGWD2', kagawad_list[1].upper())
    replace_text_version_1(doc, 'BKGWD3', kagawad_list[2].upper())
    replace_text_version_1(doc, 'BKGWD4', kagawad_list[3].upper())
    replace_text_version_1(doc, 'BKGWD5', kagawad_list[4].upper())
    replace_text_version_1(doc, 'BKGWD6', kagawad_list[5].upper())
    replace_text_version_1(doc, 'BKGWD7', kagawad_list[6].upper())
    replace_text_version_1(doc, 'BSCRTRY', get_barangay_secretary().upper())
    replace_text_version_1(doc, 'BTRSRR', get_barangay_treasurer().upper())
    replace_text_version_1(doc, 'BSKCHRMN', get_barangay_sk().upper())
    replace_text_version_1(doc, '{{BLPNMMBR1}}', lupon_list[0].upper())
    replace_text_version_1(doc, 'BLPNMMBR2', lupon_list[1].upper())
    replace_text_version_1(doc, 'BLPNMMBR3', lupon_list[2].upper())
    replace_text_version_1(doc, 'BLPNMMBR4', lupon_list[3].upper())
    replace_text_version_1(doc, 'BLPNMMBR5', lupon_list[4].upper())
    replace_text_version_1(doc, 'BLPNMMBR6', lupon_list[5].upper())
    replace_text_version_1(doc, 'BLPNMMBR7', lupon_list[6].upper())
    replace_text_version_1(doc, 'BLPNMMBR8', lupon_list[7].upper())
    replace_text_version_1(doc, 'BLPNMMBR9', lupon_list[8].upper())
    replace_text_version_1(doc, 'BLPNMMBR01', lupon_list[9].upper())
    replace_text_version_1(doc, 'YEARTODAY', str_yearToday)
    replace_text_version_1(doc, '{{YEARTODAY}}', str_yearToday)

    document_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Officials and Lupon List")
    pdf_save_path = os.path.expanduser(f"~/Desktop/{APP_NAME}/Officials and Lupon List")
    document_name = f"Barangay Poblacion 8 Officials.docx"
    pdf_name = f"Barangay Poblacion 8 Officials.pdf"
    document_path = os.path.join(document_save_path, document_name)
    pdf_path = os.path.join(pdf_save_path, pdf_name)

    os.makedirs(document_save_path, exist_ok=True)

    absolute_document_path = os.path.abspath(document_path)
    absolute_pdf_path = os.path.abspath(pdf_path)
    absolute_pdf_path = get_unique_filename(absolute_pdf_path)

    # Save the document
    doc.save(absolute_document_path)
    convert_pdf(absolute_document_path, absolute_pdf_path)
    os.remove(absolute_document_path)


# Create Directories for first time install
def do_requirements():
    dir_direct = os.path.expanduser("~\\Desktop")

    # Check if PROB8 directory exists
    dir_system_app = dir_direct + f"\\{APP_NAME}"
    if not os.path.exists(dir_system_app):
        os.makedirs(dir_system_app)

    dir_profiling_docs = dir_system_app + f"\\Residents Profiling Documents"
    if not os.path.exists(dir_profiling_docs):
        os.makedirs(dir_profiling_docs)

    dir_officials_list = dir_system_app + f"\\Officials and Lupon List"
    if not os.path.exists(dir_officials_list):
        os.makedirs(dir_officials_list)

    # Check if Certificate of Low Income directory exists
    dir_lowIncome = dir_profiling_docs + "\\Certificate of Low Income"
    if not os.path.exists(dir_lowIncome):
        os.makedirs(dir_lowIncome)

    # Check if Barangay Clearance Directory exists
    dir_clearance = dir_profiling_docs + "\\Barangay Clearance"
    if not os.path.exists(dir_clearance):
        os.makedirs(dir_clearance)

    # Check if Barangay Residency Directory exists
    dir_residency = dir_profiling_docs + "\\Barangay Residency"
    if not os.path.exists(dir_residency):
        os.makedirs(dir_residency)

    # Check if Barangay First Time Job Seekers
    dir_jobSeekers = dir_profiling_docs + "\\First Time Jobseekers"
    if not os.path.exists(dir_jobSeekers):
        os.makedirs(dir_jobSeekers)

    # Check if Certificate of Good Moral Character directory exists
    dir_goodMoral = dir_profiling_docs + "\\Certificate of Good Moral Character"
    if not os.path.exists(dir_goodMoral):
        os.makedirs(dir_goodMoral)
        
    dir_db_backup = dir_system_app + "\\Database Backup"
    if not os.path.exists(dir_db_backup):
        os.makedirs(dir_db_backup)

    dir_auto_db_backup = dir_db_backup + "\\Scheduled Backup"
    if not os.path.exists(dir_auto_db_backup):
        os.makedirs(dir_auto_db_backup)

    dir_blotter = dir_system_app + "\\Blotter Documents"
    if not os.path.exists(dir_blotter):
        os.makedirs(dir_blotter)

    dir_reports = dir_profiling_docs + "\\Reports"
    if not os.path.exists(dir_reports):
        os.makedirs(dir_reports)
